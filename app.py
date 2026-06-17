from flask import Flask, render_template, request, jsonify, session
from google import genai
import os
import json
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from src.database import init_db, register_user, login_user, create_conversation, get_user_conversations, get_conversation_messages, save_message, delete_conversation, get_user_profile, update_user_profile, change_password, get_db
from src.rag_system import RAGSystem
import PyPDF2
import docx

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'your-secret-key-change-in-production')
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize database
init_db()

# Initialize Gemini client with API key from .env
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
GEMINI_MODEL = os.getenv('GEMINI_MODEL', 'gemini-2.5-flash')

if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY not found in .env file")

client = genai.Client(api_key=GEMINI_API_KEY)
rag_system = RAGSystem(GEMINI_API_KEY, GEMINI_MODEL)

# Load simple responses
def load_simple_responses():
    """Load simple Q&A responses from JSON file"""
    try:
        with open(os.path.join(os.path.dirname(__file__), 'src', 'simple_responses.json'), 'r') as f:
            return json.load(f)['simple_responses']
    except Exception as e:
        print(f"Error loading simple responses: {e}")
        return {}

simple_responses = load_simple_responses()

def get_simple_response(message):
    """Check if message matches any simple response"""
    message_lower = message.lower().strip()
    return simple_responses.get(message_lower)

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_file_content(filepath):
    """Read content from uploaded file using custom algorithm"""
    try:
        if filepath.endswith('.txt'):
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
                return preprocess_text(content)
        elif filepath.endswith('.pdf'):
            return extract_pdf_content(filepath)
        elif filepath.endswith('.docx'):
            return extract_docx_content(filepath)
        elif filepath.endswith('.doc'):
            return "Old .doc format not supported. Please convert to .docx."
        else:
            return "Unsupported file format"
    except Exception as e:
        return f"Error reading file: {str(e)}"

def preprocess_text(text):
    """Preprocess text before sending to LLM"""
    # Remove extra whitespace
    text = ' '.join(text.split())
    # Remove special characters that might cause issues
    text = text.replace('\x00', '').replace('\r\n', '\n')
    return text

def extract_pdf_content(filepath):
    """Extract text from PDF using PyPDF2"""
    try:
        text = ""
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        return preprocess_text(text)
    except Exception as e:
        return f"Error extracting PDF content: {str(e)}"

def extract_docx_content(filepath):
    """Extract text from DOCX using python-docx"""
    try:
        doc = docx.Document(filepath)
        text = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
        
        return preprocess_text(text)
    except Exception as e:
        return f"Error extracting DOCX content: {str(e)}"


@app.route('/')
def index():
    """Render the main page"""
    return render_template('index.html')


# ============= AUTHENTICATION ENDPOINTS =============

@app.route('/api/register', methods=['POST'])
def register():
    """Register a new user"""
    try:
        data = request.json
        name = data.get('name', '').strip()
        email = data.get('email', '').strip()
        password = data.get('password', '').strip()
        
        if not email or not password:
            return jsonify({'success': False, 'error': 'Email and password required'}), 400
        
        if len(password) < 6:
            return jsonify({'success': False, 'error': 'Password must be at least 6 characters'}), 400
        
        # Use email as username
        username = email
        
        if register_user(username, password, name, email):
            session['user_id'] = login_user(username, password)
            session['username'] = username
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'error': 'Email already exists'}), 400
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/login', methods=['POST'])
def login():
    """Login user"""
    try:
        data = request.json
        username = data.get('username', '').strip()
        password = data.get('password', '').strip()
        
        if not username or not password:
            return jsonify({'success': False, 'error': 'Username and password required'}), 400
        
        user_id = login_user(username, password)
        
        if user_id:
            session['user_id'] = user_id
            session['username'] = username
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'error': 'Invalid credentials'}), 401
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/logout', methods=['POST'])
def logout():
    """Logout user"""
    session.clear()
    return jsonify({'success': True})


@app.route('/api/user', methods=['GET'])
def get_user():
    """Get current user info"""
    if 'user_id' in session:
        profile = get_user_profile(session['user_id'])
        return jsonify({
            'success': True,
            'user_id': session['user_id'],
            'username': session['username'],
            'profile': profile
        })
    return jsonify({'success': False, 'error': 'Not logged in'}), 401


# ============= CONVERSATION ENDPOINTS =============

@app.route('/api/conversations', methods=['GET'])
def get_conversations():
    """Get all conversations for current user"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    try:
        conversations = get_user_conversations(session['user_id'])
        return jsonify({'success': True, 'conversations': conversations})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/conversations/new', methods=['POST'])
def new_conversation():
    """Create a new conversation"""
    try:
        data = request.json
        title = data.get('title') if data else None
        
        if 'user_id' in session:
            conversation_id = create_conversation(session['user_id'], title)
            return jsonify({'success': True, 'conversation_id': conversation_id})
        else:
            # For non-logged-in users, create a temporary conversation with user_id = 0
            conversation_id = create_conversation(0, title)
            return jsonify({'success': True, 'conversation_id': conversation_id})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/conversations/<int:conversation_id>/messages', methods=['GET'])
def get_messages(conversation_id):
    """Get all messages in a conversation"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    try:
        # Verify conversation ownership
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute('SELECT user_id FROM conversations WHERE id = ?', (conversation_id,))
        result = cursor.fetchone()
        conn.close()
        
        if not result or result[0] != session['user_id']:
            return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        messages = get_conversation_messages(conversation_id)
        return jsonify({'success': True, 'messages': messages})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/conversations/<int:conversation_id>/delete', methods=['POST'])
def delete_conv(conversation_id):
    """Delete a conversation"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    try:
        # Verify conversation ownership
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute('SELECT user_id FROM conversations WHERE id = ?', (conversation_id,))
        result = cursor.fetchone()
        conn.close()
        
        if not result or result[0] != session['user_id']:
            return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        delete_conversation(conversation_id)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============= CHAT ENDPOINTS =============

@app.route('/api/chat', methods=['POST'])
def chat():
    """API endpoint to handle chat requests with RAG"""
    try:
        data = request.json
        user_message = data.get('message', '').strip()
        conversation_id = data.get('conversation_id')
        
        if not user_message:
            return jsonify({'error': 'Message cannot be empty'}), 400
        
        if not conversation_id:
            return jsonify({'error': 'Conversation ID required'}), 400
        
        # Verify conversation ownership if user is logged in
        if 'user_id' in session:
            conn = get_db()
            cursor = conn.cursor()
            cursor.execute('SELECT user_id FROM conversations WHERE id = ?', (conversation_id,))
            result = cursor.fetchone()
            conn.close()
            
            if result and result[0] != session['user_id'] and result[0] != 0:
                return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        # Check for simple response first
        simple_response = get_simple_response(user_message)
        if simple_response:
            response_text = simple_response
        else:
            # Get conversation history for context
            previous_messages = get_conversation_messages(conversation_id)
            
            # Generate response using RAG system
            response_text = rag_system.generate_response_with_rag(user_message, previous_messages)
        
        # Save user message
        save_message(conversation_id, 'user', user_message)
        
        # Save assistant response
        save_message(conversation_id, 'assistant', response_text)
        
        return jsonify({
            'success': True,
            'message': response_text
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/models', methods=['GET'])
def get_models():
    """Get the current model configuration"""
    return jsonify({
        'model': GEMINI_MODEL,
        'provider': 'Google Gemini'
    })


# ============= PROFILE ENDPOINTS =============

@app.route('/api/profile', methods=['GET'])
def get_profile():
    """Get user profile"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    try:
        profile = get_user_profile(session['user_id'])
        return jsonify({'success': True, 'profile': profile})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/profile/update', methods=['POST'])
def update_profile():
    """Update user profile"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    try:
        data = request.json
        name = data.get('name')
        mobile = data.get('mobile')
        dob = data.get('dob')
        
        update_user_profile(session['user_id'], name, mobile, dob)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/profile/change-password', methods=['POST'])
def change_password_endpoint():
    """Change user password"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    try:
        data = request.json
        current_password = data.get('current_password')
        new_password = data.get('new_password')
        
        if not current_password or not new_password:
            return jsonify({'success': False, 'error': 'Current and new password required'}), 400
        
        # Verify current password
        username = session['username']
        if not login_user(username, current_password):
            return jsonify({'success': False, 'error': 'Current password is incorrect'}), 400
        
        if len(new_password) < 6:
            return jsonify({'success': False, 'error': 'New password must be at least 6 characters'}), 400
        
        change_password(session['user_id'], new_password)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============= DOCUMENT ENDPOINTS =============

@app.route('/api/upload-document', methods=['POST'])
def upload_document():
    """Upload and summarize a document"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'error': 'Invalid file type. Allowed: txt, pdf, doc, docx'}), 400
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Read file content using custom algorithm
        content = read_file_content(filepath)
        
        # Delete the file after reading (no persistence needed)
        try:
            os.remove(filepath)
        except:
            pass
        
        # Generate summary using AI
        summary_prompt = f"Please provide a concise summary of the following document content:\n\n{content}\n\nSummary:"
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=summary_prompt
        )
        summary = response.text
        
        # Generate topics/labels using AI
        topics_prompt = f"Extract 3-5 key topics or labels from the following document summary. Return them as a comma-separated list:\n\n{summary}\n\nTopics:"
        topics_response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=topics_prompt
        )
        topics = topics_response.text
        
        # Return the data without saving to database
        return jsonify({
            'success': True,
            'summary': summary,
            'topics': topics,
            'filename': filename,
            'content': content
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/generate-document', methods=['POST'])
def generate_document():
    """Generate a document based on provided data"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    try:
        data = request.json
        doc_type = data.get('doc_type')
        doc_format = data.get('doc_format')
        doc_data = data.get('doc_data')
        
        if not doc_type or not doc_format or not doc_data:
            return jsonify({'success': False, 'error': 'All fields are required'}), 400
        
        # Generate document using AI
        prompt = f"""Generate a {doc_type} document in {doc_format} format based on the following data/instructions:

{doc_data}

Please provide the complete document content."""
        
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt
        )
        generated_content = response.text
        
        return jsonify({
            'success': True,
            'content': generated_content,
            'format': doc_format
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/document-chat', methods=['POST'])
def document_chat():
    """Chat with document using RAG"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    try:
        data = request.json
        user_message = data.get('message', '').strip()
        document_content = data.get('document_content', '')
        document_summary = data.get('document_summary', '')
        
        if not user_message:
            return jsonify({'error': 'Message cannot be empty'}), 400
        
        # Use RAG with document context, allowing general knowledge for terms in document
        prompt = f"""Based on the following document information, answer the user's question:

Document Summary:
{document_summary}

Document Content:
{document_content}

User Question: {user_message}

IMPORTANT INSTRUCTIONS:
1. First, check if the question relates to terms, concepts, or topics that are explicitly mentioned in the document.
2. If the question asks about the meaning or explanation of a term/concept that appears in the document, you may use your general knowledge to explain it, but only if that term is actually present in the document content.
3. If the question asks about something completely unrelated to the document or terms not mentioned in the document, politely state that this topic is outside the scope of the document.
4. Provide accurate, helpful answers based on the document content first, then supplement with general knowledge only for terms explicitly mentioned in the document.

Please provide a helpful answer following these guidelines."""
        
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt
        )
        
        return jsonify({
            'success': True,
            'response': response.text
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


if __name__ == '__main__':
    debug = os.getenv('DEBUG', 'True').lower() == 'true'
    app.run(debug=debug, host='0.0.0.0', port=5000)