from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from utils.doc_builders import _add_hyperlink

def _html_to_docx_bytes(html, title='Document'):
    doc = Document()
    soup = BeautifulSoup(html or '', 'html.parser')
    body = soup.body or soup

    def add_runs(p, node, bold=False, italic=False, underline=False):
        if hasattr(node, 'name') and node.name is None: return
        if not hasattr(node, 'name') or node.name is None:
            txt = str(node)
            if txt:
                r = p.add_run(txt)
                r.bold, r.italic, r.underline = bold, italic, underline
                r.font.color.rgb = RGBColor(0, 0, 0)
            return
        name = node.name.lower() if node.name else ''
        new_bold = bold or name in ('b', 'strong')
        new_italic = italic or name in ('i', 'em')
        new_underline = underline or name == 'u'
        if name == 'br':
            p.add_run().add_break(); return
        if name == 'a':
            text = node.get_text()
            _add_hyperlink(p, node.get('href') or text, text, bold=new_bold)
            return
        for child in node.children:
            if hasattr(child, 'children') and getattr(child, 'name', None):
                add_runs(p, child, new_bold, new_italic, new_underline)
            else:
                txt = str(child)
                if txt:
                    r = p.add_run(txt)
                    r.bold, r.italic, r.underline = new_bold, new_italic, new_underline
                    r.font.color.rgb = RGBColor(0, 0, 0)

    def walk(node):
        for child in node.children:
            if not hasattr(child, 'name') or child.name is None:
                if str(child).strip(): add_runs(doc.add_paragraph(), child)
                continue
            tag = child.name.lower()
            if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                p = doc.add_heading('', level=min(int(tag[1]), 4))
                for c in child.children: add_runs(p, c)
                for r in p.runs: r.font.color.rgb = RGBColor(0, 0, 0)
            elif tag in ('p', 'div'):
                add_runs(doc.add_paragraph(), child)
            elif tag in ('ul', 'ol'):
                style = 'List Bullet' if tag == 'ul' else 'List Number'
                for li in child.find_all('li', recursive=False): add_runs(doc.add_paragraph(style=style), li)
            elif tag == 'blockquote':
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(24)
                for c in child.children: add_runs(p, c, italic=True)
            elif tag == 'pre':
                doc.add_paragraph().add_run(child.get_text()).font.name = 'Consolas'
            elif tag == 'hr':
                doc.add_paragraph('—' * 30)
            elif tag == 'table':
                rows = child.find_all('tr')
                if not rows: continue
                cols = max(len(r.find_all(['td', 'th'])) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=cols); tbl.style = 'Table Grid'
                for ri, tr in enumerate(rows):
                    for ci, cell in enumerate(tr.find_all(['td', 'th'])): tbl.rows[ri].cells[ci].text = cell.get_text()
            else:
                add_runs(doc.add_paragraph(), child)
    walk(body)
    buf = BytesIO(); doc.save(buf); buf.seek(0); return buf

def _html_to_pdf_bytes(html, title='Document'):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=1.8*cm, rightMargin=1.8*cm, topMargin=2*cm, bottomMargin=2*cm)
    ss = getSampleStyleSheet()
    BLACK = colors.HexColor('#000000')
    body = ParagraphStyle('rb', parent=ss['BodyText'], textColor=BLACK, fontSize=11, leading=15, spaceAfter=6)
    h1 = ParagraphStyle('rh1', parent=ss['Heading1'], textColor=BLACK, fontSize=18, spaceAfter=8)
    h2 = ParagraphStyle('rh2', parent=ss['Heading2'], textColor=BLACK, fontSize=15, spaceAfter=6)
    h3 = ParagraphStyle('rh3', parent=ss['Heading3'], textColor=BLACK, fontSize=13, spaceAfter=4)
    soup = BeautifulSoup(html or '', 'html.parser')
    body_node = soup.body or soup
    story = []

    def inline_html(node):
        s = ''
        for c in node.children:
            if not hasattr(c, 'name') or c.name is None:
                s += str(c).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            else:
                tag = c.name.lower()
                inner = inline_html(c)
                if tag in ('b', 'strong'): s += f'<b>{inner}</b>'
                elif tag in ('i', 'em'): s += f'<i>{inner}</i>'
                elif tag == 'u': s += f'<u>{inner}</u>'
                elif tag == 'a': s += f'<a href="{c.get("href") or inner}"><font color="#0563C1"><u>{inner}</u></font></a>'
                elif tag == 'br': s += '<br/>'
                else: s += inner
        return s

    for child in body_node.children:
        if not hasattr(child, 'name') or child.name is None:
            if str(child).strip(): story.append(Paragraph(str(child).strip(), body))
            continue
        tag = child.name.lower()
        if tag in ('h1', 'h2', 'h3', 'h4'):
            story.append(Paragraph(inline_html(child), {'h1': h1, 'h2': h2}.get(tag, h3)))
        elif tag in ('p', 'div'):
            story.append(Paragraph(inline_html(child), body))
        elif tag == 'hr':
            story.append(HRFlowable(width='100%', thickness=0.5, color=BLACK, spaceBefore=4, spaceAfter=6))
        elif tag in ('ul', 'ol'):
            for i, li in enumerate(child.find_all('li', recursive=False), 1):
                story.append(Paragraph(('• ' if tag == 'ul' else f'{i}. ') + inline_html(li), body))
        elif tag == 'blockquote':
            story.append(Paragraph('<i>' + inline_html(child) + '</i>', body))
        elif tag == 'pre':
            story.append(Paragraph('<font name="Courier">' + child.get_text().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>') + '</font>', body))
        elif tag == 'table':
            rows = child.find_all('tr')
            data_rows = [[cell.get_text() for cell in tr.find_all(['td', 'th'])] for tr in rows]
            if data_rows:
                t = Table(data_rows); t.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.4, colors.grey)]))
                story.append(t); story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(inline_html(child), body))
    if not story: story.append(Paragraph(' ', body))
    doc.build(story); buf.seek(0); return buf