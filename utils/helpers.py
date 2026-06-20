import os
import re
import requests
import json
import PyPDF2
import docx

RECAPTCHA_SECRET_KEY = os.getenv('RECAPTCHA_SECRET_KEY', '')
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

def _load_simple_responses():
    try:
        path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'backend/src', 'simple_responses.json')
        with open(path) as f:
            return json.load(f)['simple_responses']
    except Exception:
        return {}

simple_responses = _load_simple_responses()

def verify_recaptcha(token):
    if (os.getenv('DEBUG', 'False').lower() == 'true') and (token in (None, '', 'dummy', 'dev-bypass')):
        return True
    if not token or not RECAPTCHA_SECRET_KEY:
        return not RECAPTCHA_SECRET_KEY
    try:
        r = requests.post(
            'https://www.google.com/recaptcha/api/siteverify',
            data={'secret': RECAPTCHA_SECRET_KEY, 'response': token}, timeout=5,
        )
        return r.json().get('success', False)
    except Exception:
        return False

def allowed_file(fn):
    return '.' in fn and fn.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def preprocess_text(t):
    return ' '.join((t or '').split()).replace('\x00', '')

def read_file_content(filepath):
    try:
        if filepath.endswith('.txt'):
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return preprocess_text(f.read())
        if filepath.endswith('.pdf'):
            text = ''
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for p in reader.pages:
                    text += (p.extract_text() or '') + '\n'
            return preprocess_text(text)
        if filepath.endswith('.docx'):
            d = docx.Document(filepath)
            text = '\n'.join(p.text for p in d.paragraphs)
            for table in d.tables:
                for row in table.rows:
                    text += '\n' + ' | '.join(c.text for c in row.cells)
            return preprocess_text(text)
        return 'Unsupported file format'
    except Exception as e:
        return f'Error reading file: {e}'

def safe_name(s):
    s = re.sub(r'[^A-Za-z0-9_\- ]+', '', s or '').strip().replace(' ', '_') or 'document'
    return s[:60]